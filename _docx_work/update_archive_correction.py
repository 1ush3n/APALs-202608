from pathlib import Path

from docx import Document


root = Path(r"D:\APAL-Dynamic-v4")
target = next(root.joinpath("docs").glob("*.docx"))
output = root / "_docx_work" / "archive_with_correction.docx"
doc = Document(target)
marker = "附录 G 结果身份更正：旧 operation-only/operation+station 重调度结果归档边界"
if any(p.text.strip() == marker for p in doc.paragraphs):
    raise SystemExit("目标文档已包含附录 G，停止以避免重复追加。")

doc.add_page_break()
doc.add_heading(marker, level=1)
doc.add_paragraph(
    "口径更正：前述 r5 扰动强度排名中的 HB-GAT-PPO operation-only 与 HB-GAT-PPO operation+station，是旧的重调度动作范围结果，不是本次初始调度的 EarliestAvailabilityActionCompleter（EA）结果。它们与 operation_only_ea、operation_station_ea 的实验身份不同，不能用来证明最短等待 EA 的性能。"
)
doc.add_heading("G.1 归档与正式对比边界", level=2)
for text in [
    "旧 operation-only 与 operation+station 的 r5 原始逐场景数据、汇总表和审计文件完整保留在 results/r5_new_weight_comparison_20260827/，并标记为“已存档，不参与正式对比”；不删除、不移动、不重命名原始结果。",
    "results/revalidation_20260829/ 中的 operation_only_ea 与 operation_station_ea 是初始调度 EA 诊断复验，使用旧 checkpoint、单 seed 和 temperature=0；该目录没有 low/medium/high 扰动强度的 r5 正式重调度结果。",
    "因此，当前正式重调度对比不得继续使用上述两个旧动作范围方法的排名；EA 的正式重调度排名须等待对应 EA 变体完成正式训练，并按相同 r5 场景协议重新汇总。",
]:
    doc.add_paragraph(text, style="List Bullet")
doc.add_paragraph(
    "本更正不否定旧结果的历史复核价值，只修正其方法身份和正式对比资格，避免将旧重调度动作范围误读为当前初始调度 EA。"
)
doc.save(output)
print(output)
